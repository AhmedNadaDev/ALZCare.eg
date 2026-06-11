# -*- coding: utf-8 -*-
"""
Build the IEEE-formatted graduation paper (.docx) for the ALZCare
speech-based mood detection system + RAG clinical assistant.

Two-column IEEE layout, spanning title block, abstract/keywords, numbered
sections, professional tables, embedded evaluation figures with captions,
and an IEEE-style numbered reference list.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

BASE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(BASE, "figs")
MOOD = os.path.join(BASE, "evalatuon mod")
RAG = os.path.join(BASE, "evaluation rag")
OUT = os.path.join(BASE, "output")
os.makedirs(OUT, exist_ok=True)

COL_W = 3.42          # inches, usable single-column width
BODY_FONT = "Times New Roman"
BODY_PT = 9.5

# image files
F_SYS = os.path.join(FIG, "fig_system_arch.png")
F_MODEL = os.path.join(FIG, "fig_model_arch.png")
F_RAG = os.path.join(FIG, "fig_rag_arch.png")
M1 = os.path.join(MOOD, "photo_1_2026-06-11_17-21-36.jpg")   # confusion matrix
M2 = os.path.join(MOOD, "photo_2_2026-06-11_17-21-36.jpg")   # val macro-F1
M3 = os.path.join(MOOD, "photo_3_2026-06-11_17-21-36.jpg")   # calibration
R1 = os.path.join(RAG, "photo_1_2026-06-11_17-20-52.jpg")    # cumulative similarity curve
R2 = os.path.join(RAG, "photo_2_2026-06-11_17-20-52.jpg")    # retrieval summary bars
R3 = os.path.join(RAG, "photo_3_2026-06-11_17-20-52.jpg")    # per-sample similarity
R4 = os.path.join(RAG, "photo_4_2026-06-11_17-20-52.jpg")    # violin distribution
R5 = os.path.join(RAG, "photo_5_2026-06-11_17-20-52.jpg")    # radar
R6 = os.path.join(RAG, "photo_6_2026-06-11_17-20-52.jpg")    # cumulative stability


# ----------------------------------------------------------------------------- helpers
def set_cols(section, n, space_twips=240):
    cols = section._sectPr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), str(n))
    cols.set(qn("w:space"), str(space_twips))


def set_margins(section, top=0.7, bottom=0.85, left=0.62, right=0.62):
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)
    section.gutter = Inches(0)


def _no_space(pf):
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def title_block(doc, title, authors, affil_lines):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_space(p.paragraph_format)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(title)
    r.font.name = BODY_FONT
    r.font.size = Pt(20)
    r.bold = True

    pa = doc.add_paragraph()
    pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_space(pa.paragraph_format)
    pa.paragraph_format.space_after = Pt(2)
    ra = pa.add_run(authors)
    ra.font.name = BODY_FONT
    ra.font.size = Pt(11)

    for i, line in enumerate(affil_lines):
        pf = doc.add_paragraph()
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _no_space(pf.paragraph_format)
        if i == len(affil_lines) - 1:
            pf.paragraph_format.space_after = Pt(8)
        rf = pf.add_run(line)
        rf.font.name = BODY_FONT
        rf.font.size = Pt(10)
        rf.italic = True


def lead_para(doc, label, body):
    """Abstract / Index Terms — first run bold-italic label, rest italic, justified, no indent."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _no_space(p.paragraph_format)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Inches(0)
    rl = p.add_run(label)
    rl.font.name = BODY_FONT
    rl.font.size = Pt(BODY_PT)
    rl.bold = True
    rl.italic = True
    rb = p.add_run(body)
    rb.font.name = BODY_FONT
    rb.font.size = Pt(BODY_PT)
    rb.italic = True
    return p


def h1(doc, number, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_space(p.paragraph_format)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run((str(number) + ". " + text) if number else text)
    r.font.name = BODY_FONT
    r.font.size = Pt(BODY_PT)
    r.font.small_caps = True
    return p


def h2(doc, letter, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _no_space(p.paragraph_format)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{letter}. {text}")
    r.font.name = BODY_FONT
    r.font.size = Pt(BODY_PT)
    r.italic = True
    return p


def body(doc, text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _no_space(p.paragraph_format)
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.16)
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(BODY_PT)
    return p


def bullet(doc, label, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _no_space(p.paragraph_format)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    rb = p.add_run(("• " + label) if label else "• ")
    rb.font.name = BODY_FONT
    rb.font.size = Pt(BODY_PT)
    rb.bold = bool(label)
    rt = p.add_run(text)
    rt.font.name = BODY_FONT
    rt.font.size = Pt(BODY_PT)
    return p


def add_figure(doc, path, caption, width=COL_W):
    im = Image.open(path)
    w, h = im.size
    disp_w = min(width, COL_W)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_space(p.paragraph_format)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run()
    run.add_picture(path, width=Inches(disp_w))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _no_space(c.paragraph_format)
    c.paragraph_format.space_after = Pt(6)
    rb = c.add_run(caption[0])
    rb.font.name = BODY_FONT
    rb.font.size = Pt(8)
    rb.bold = False
    rt = c.add_run(" " + caption[1])
    rt.font.name = BODY_FONT
    rt.font.size = Pt(8)


def _set_cell(cell, text, bold=False, size=8, align="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    _no_space(p.paragraph_format)
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(size)
    r.bold = bold


def _set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "insideH"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    for edge in ("left", "right", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    tblPr.append(borders)


def add_table(doc, number, title, headers, rows, col_w=None, size=8):
    # caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_space(cap.paragraph_format)
    cap.paragraph_format.space_before = Pt(6)
    rc = cap.add_run(f"TABLE {number}")
    rc.font.name = BODY_FONT
    rc.font.size = Pt(8)
    cap2 = doc.add_paragraph()
    cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_space(cap2.paragraph_format)
    cap2.paragraph_format.space_after = Pt(2)
    rc2 = cap2.add_run(title)
    rc2.font.name = BODY_FONT
    rc2.font.size = Pt(8)
    rc2.font.small_caps = True

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, htxt in enumerate(headers):
        _set_cell(hdr[i], htxt, bold=True, size=size, align="center")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            _set_cell(cells[i], val, bold=False, size=size,
                      align="left" if i == 0 else "center")
    _set_table_borders(table)
    if col_w:
        for i, w in enumerate(col_w):
            for r in table.rows:
                r.cells[i].width = Inches(w)
    # spacing after table
    sp = doc.add_paragraph()
    _no_space(sp.paragraph_format)
    sp.paragraph_format.space_after = Pt(4)
    return table


def add_ref(doc, idx, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _no_space(p.paragraph_format)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    r = p.add_run(f"[{idx}] {text}")
    r.font.name = BODY_FONT
    r.font.size = Pt(8)


# ============================================================================= BUILD
doc = Document()

# base style
style = doc.styles["Normal"]
style.font.name = BODY_FONT
style.font.size = Pt(BODY_PT)

# section 0 (single column, title block)
sec0 = doc.sections[0]
set_margins(sec0)
set_cols(sec0, 1)

TITLE = ("ALZCare: A Speech-Based Multi-Task Mood Detection System with a "
         "Retrieval-Augmented Clinical Assistant for Alzheimer’s Care")
AUTHORS = "Ahmed Nada, Omar Ibrahim, Ahmed Fathy, Aya Abdelmagid, Negma, Rokia Islam"
AFFIL = ["Department of Computer Engineering",
         "Alamein International University, Alamein, Egypt"]
title_block(doc, TITLE, AUTHORS, AFFIL)

# continuous break -> two columns for the body
doc.add_section(WD_SECTION.CONTINUOUS)
sec1 = doc.sections[1]
set_margins(sec1)
set_cols(sec1, 2, space_twips=288)

# --------------------------------------------------------------------------- ABSTRACT
ABSTRACT = (
    "Mood and arousal are clinically meaningful signals for monitoring patients living with "
    "Alzheimer’s disease, yet they are difficult to capture continuously and objectively in "
    "everyday care. We present ALZCare, an integrated clinical decision-support platform that "
    "estimates a patient’s mood directly from short voice recordings and exposes the results to "
    "clinicians and family caregivers through role-based dashboards. The mood estimator is a "
    "custom multi-task neural network built on a self-supervised WavLM-Base-Plus speech encoder, "
    "an attentive masked-pooling layer, and two prediction heads that jointly produce a six-class "
    "mood label (Calm, Neutral, Content, Anxious, Agitated, Low) and a coarse low/high arousal "
    "signal. The model is trained on four merged acted-emotion corpora under class-balanced focal "
    "loss with mild, domain-aware audio augmentation, and its probabilities are calibrated by "
    "post-hoc temperature scaling. On a speaker-disjoint test split the system reaches "
    "approximately 0.74–0.77 six-class accuracy, 0.73–0.76 macro-F1, and 0.86–0.88 "
    "low/high arousal accuracy, with a peak validation macro-F1 of 0.82 and an expected "
    "calibration error reduced from 0.041 to 0.039 after calibration. Alongside the acoustic "
    "model, a retrieval-augmented generation (RAG) assistant grounds a large language model in a "
    "curated Alzheimer’s knowledge base and the patient’s own medical record; on a held-out "
    "query set it attains 0.95 retrieval precision, 0.95 mean reciprocal rank, 0.97 normalized "
    "discounted cumulative gain, 1.00 routing accuracy, and a 0.96 mean answer–reference "
    "semantic similarity. We describe the end-to-end MERN architecture, the inference pipelines, "
    "the evaluation methodology, and the strengths and limitations of the deployed prototype."
)
lead_para(doc, "Abstract—", ABSTRACT)
KEYWORDS = (
    "speech emotion recognition, mood detection, WavLM, self-supervised learning, multi-task "
    "learning, model calibration, retrieval-augmented generation, large language models, "
    "Alzheimer’s disease, clinical decision support."
)
lead_para(doc, "Index Terms—", KEYWORDS)

# --------------------------------------------------------------------------- I. INTRO
h1(doc, "I", "Introduction")
body(doc,
    "Alzheimer’s disease and related dementias progressively impair memory, communication, "
    "and emotional regulation. Among the most burdensome aspects of the disease for both patients "
    "and caregivers are the behavioural and psychological symptoms of dementia—agitation, "
    "anxiety, apathy, and depressed mood—which fluctuate over time and are strongly predictive "
    "of caregiver stress and institutionalisation. Continuous, objective monitoring of a "
    "patient’s affective state would allow clinicians and families to detect distress earlier "
    "and to respond before episodes escalate. In practice, however, mood is assessed sporadically "
    "through subjective questionnaires and brief clinical visits, leaving most of a patient’s "
    "daily emotional trajectory unobserved.", indent=False)
body(doc,
    "Speech is an attractive modality for this problem. It is non-invasive, can be captured with "
    "commodity devices, and carries rich paralinguistic information—prosody, energy, and "
    "timbre—that correlates with affect independently of the words spoken. Recent advances in "
    "self-supervised speech representation learning have made it possible to build accurate "
    "acoustic models from limited labelled data, and parallel advances in large language models "
    "(LLMs) and retrieval-augmented generation (RAG) have made it possible to deliver grounded, "
    "trustworthy clinical information to non-expert users. The challenge is to combine these "
    "capabilities into a single, safe, and usable system.")
body(doc,
    "This paper presents ALZCare, a deployed full-stack prototype that does exactly this. A "
    "patient records a short voice clip; a speech mood model estimates the patient’s mood and "
    "arousal; the result is stored, surfaced to the patient’s doctor and family in real time, "
    "and flagged when it indicates possible distress. A complementary RAG assistant answers "
    "clinical and caregiving questions, grounded strictly in the patient’s record and a curated "
    "medical knowledge base so that it cannot fabricate patient facts. The system is explicitly a "
    "decision-support tool—it produces calibrated probabilities and abstains when uncertain, "
    "rather than issuing diagnoses.")
body(doc, "The contributions of this work are as follows:")
bullet(doc, "A multi-task speech mood model.",
    " a WavLM-Base-Plus backbone with attentive masked pooling and joint mood/arousal heads, "
    "trained with class-balanced focal loss and calibrated by temperature scaling, that produces "
    "six-class mood estimates together with a more reliable coarse arousal signal.")
bullet(doc, "An integrated, role-aware care platform.",
    " a MERN-stack application that captures voice check-ins in the browser, serves model "
    "inference through Python micro-services, persists calibrated results, and broadcasts distress "
    "alerts to clinicians and caregivers over WebSockets.")
bullet(doc, "A clinically-safe RAG assistant.",
    " a semantically-routed retrieval pipeline that grounds an LLM in a FAISS knowledge base and "
    "the MongoDB patient record, with explicit anti-hallucination guardrails and role-adaptive "
    "(clinician vs. family) responses.")
bullet(doc, "An empirical evaluation.",
    " a quantitative characterisation of both subsystems—classification accuracy, macro-F1, "
    "calibration, per-class confusion for the mood model, and retrieval precision, MRR, NDCG, "
    "routing accuracy, and answer semantic similarity for the RAG assistant.")
body(doc,
    "The remainder of the paper is organised as follows. Section II reviews related work. "
    "Section III describes the system architecture. Section IV details the methodology of both "
    "AI subsystems. Section V reports the experimental evaluation. Section VI discusses strengths "
    "and limitations, and Section VII concludes with directions for future work.")

# --------------------------------------------------------------------------- II. RELATED
h1(doc, "II", "Related Work")
h2(doc, "A", "Speech Emotion Recognition")
body(doc,
    "Speech emotion recognition (SER) has a long history rooted in hand-crafted acoustic "
    "descriptors—pitch, energy, mel-frequency cepstral coefficients (MFCCs), and their "
    "statistical functionals—fed to classical classifiers. The field shifted decisively with "
    "self-supervised speech representation models such as wav2vec 2.0 [1], HuBERT [2], and "
    "WavLM [3], which are pre-trained on large unlabelled corpora and then fine-tuned for "
    "downstream tasks. WavLM in particular augments masked-prediction pre-training with denoising "
    "and overlapped-speech simulation, yielding representations that are robust to noise and "
    "speaker variation; it is a strong performer across the SUPERB benchmark [4], including "
    "emotion recognition. ALZCare adopts WavLM-Base-Plus as its acoustic backbone for these "
    "reasons.", indent=False)
body(doc,
    "A recurring difficulty in SER is that prosody encodes arousal far more strongly than "
    "valence: high-arousal states such as happiness and anger are acoustically similar, as are "
    "low-arousal negative states such as sadness and fear. Much of the residual error on acted "
    "emotion data is therefore irreducible label ambiguity rather than model underfitting. This "
    "motivates our multi-task design, in which a coarse arousal axis is predicted jointly with "
    "fine mood and treated as the more reliable clinical signal.")
h2(doc, "B", "Emotion Datasets and Robust Training")
body(doc,
    "Publicly available emotional-speech corpora are predominantly acted. RAVDESS [5], "
    "CREMA-D [6], TESS, and SAVEE provide categorical emotion labels under controlled recording "
    "conditions; they differ markedly in speaker count and balance. Training a robust classifier "
    "across such heterogeneous, imbalanced sources requires care: class-balanced losses based on "
    "the effective number of samples [7] and focal loss [8] mitigate imbalance and hard-example "
    "dominance, while SpecAugment [9] and waveform-level augmentation improve generalisation. "
    "Because neural classifiers are typically over-confident, post-hoc temperature scaling [10] is "
    "a standard, effective calibration step. ALZCare combines all of these techniques.", indent=False)
h2(doc, "C", "Retrieval-Augmented Generation for Clinical Assistance")
body(doc,
    "Large language models are fluent but prone to hallucination, which is unacceptable in a "
    "clinical setting. Retrieval-augmented generation [11] addresses this by conditioning "
    "generation on documents retrieved from an external store, improving factuality and "
    "traceability. Practical RAG systems pair a dense sentence encoder such as Sentence-BERT [12] "
    "or MiniLM [13] with an approximate nearest-neighbour index such as FAISS [14], and an "
    "instruction-tuned generator such as Llama 3 [15]. ALZCare’s assistant follows this recipe "
    "but adds a semantic router and strict grounding rules that make the patient’s database "
    "record—not the language model—the sole source of truth for patient-specific facts.", indent=False)

# --------------------------------------------------------------------------- III. ARCH
h1(doc, "III", "System Architecture")
body(doc,
    "ALZCare is a full-stack MERN application (MongoDB, Express, React, Node.js) organised around "
    "three role-based clients—patient, doctor, and family caregiver—and two dedicated "
    "Python AI micro-services. Fig. 1 shows the high-level architecture. The design separates the "
    "stateless application tier (which handles authentication, persistence, and real-time "
    "messaging) from the compute-heavy AI tier (which runs the speech model and the RAG pipeline), "
    "so that model inference can be scaled, restarted, or replaced independently of the web "
    "application.", indent=False)
add_figure(doc, F_SYS,
    ("Fig. 1.", "End-to-end ALZCare architecture. Role-based React clients communicate with a "
     "Node.js/Express API over HTTPS and Socket.IO. The API persists data in MongoDB and delegates "
     "AI workloads to two Python sidecars: a FastAPI mood-detection service hosting the WavLM "
     "model, and a retrieval-augmented chatbot service. Distress alerts are pushed back to "
     "clinician and caregiver clients in real time."))
h2(doc, "A", "Application Tier")
body(doc,
    "The backend is a Node.js/Express service (port 5001) using Mongoose as the MongoDB object "
    "document mapper. Three user roles—doctor, family, and patient—are authenticated "
    "independently through dedicated middleware that validates a signed JSON Web Token carrying a "
    "role claim; tokens are issued with a seven-day lifetime. A doctor manages many patients, "
    "whereas a family account is linked one-to-one to a single patient and is further constrained "
    "by granular permissions (for example, the ability to view medications, confirm a dose, add a "
    "mood entry, or contact the doctor). The React/Vite single-page application renders a distinct "
    "dashboard per role.", indent=False)
body(doc,
    "Real-time communication uses Socket.IO over WebSockets with a polling fallback. Each patient "
    "is associated with a logical room (“patient:{id}”) that both the patient client and the "
    "linked family client join; the server pushes mood updates, scheduled check-in prompts, and "
    "abnormal-mood alerts into this room. The principal MongoDB collections include patients, "
    "doctors, families, AI-mood records, mood schedules, medications and their logs, "
    "appointments, and notifications.")
h2(doc, "B", "AI Inference Tier")
body(doc,
    "Two Python services run alongside the Node backend. The mood-detection service is a FastAPI "
    "application (default port 8001) that hosts the WavLM multi-task model and exposes a "
    "“/mood/analyze” endpoint accepting an uploaded audio file and a “/health” probe. "
    "The chatbot service (default port 8000) hosts the RAG pipeline and exposes a "
    "“/chat/ask” endpoint. The Node backend never executes PyTorch itself; it forwards the "
    "raw audio buffer or the user question to the relevant service over localhost HTTP and returns "
    "the structured result. Service URLs are configured through environment variables, and the "
    "client wraps each call with explicit timeouts (45 s for mood inference, 60 s for chat) "
    "and health checks so that a failure in the AI tier degrades gracefully rather than blocking "
    "the application.", indent=False)
h2(doc, "C", "Data Flow of a Voice Mood Check-in")
body(doc,
    "A check-in begins either on demand or when the scheduler fires a prompt. The patient client "
    "records roughly twelve seconds of audio through the browser’s getUserMedia API and, using "
    "the Web Audio API, re-samples it to 16 kHz mono and encodes it as WAV, applying a basic "
    "silence/duration sanity check before upload. The clip is sent as multipart form data to the "
    "Node API, which (a) validates the file, (b) forwards it to the mood service, (c) persists the "
    "returned mood, arousal, confidence, and top-k probabilities as an AI-mood document, (d) "
    "creates a notification for the care team if the result is flagged abnormal, (e) emits a "
    "“mood:updated” event to the patient room, and (f) returns the result to the client. "
    "Abnormal flagging is triggered for the mood classes Anxious, Agitated, and Low, or when the "
    "top confidence falls below 0.35. Doctor and family dashboards aggregate the stored records "
    "into thirty-day mood trends and summary statistics.", indent=False)

# Table I - deployment configuration
add_table(doc, "I", "Deployment and Integration Configuration",
    ["Component", "Configuration"],
    [["Backend", "Node.js / Express, port 5001"],
     ["Database", "MongoDB via Mongoose ODM"],
     ["Frontend", "React 18 + Vite, Socket.IO client"],
     ["Real-time", "Socket.IO (WebSocket + polling)"],
     ["Roles", "Doctor, Family, Patient (JWT, 7-day)"],
     ["Mood service", "FastAPI, port 8001, /mood/analyze"],
     ["Chatbot service", "RAG, port 8000, /chat/ask"],
     ["Audio upload", "≤25 MB; wav/webm/ogg/m4a/mp3/…"],
     ["Mood timeout", "45 s (inference), 5 s (health)"]],
    col_w=[1.25, 2.0])

# --------------------------------------------------------------------------- IV. METHOD
h1(doc, "IV", "Methodology")
h2(doc, "A", "Audio Capture and Pre-processing")
body(doc,
    "All audio is reduced to a canonical form before inference: a single-channel waveform sampled "
    "at 16 kHz, the native rate of the WavLM feature extractor. Browser-recorded audio is "
    "converted client-side, and the mood service additionally decodes and resamples any accepted "
    "container server-side. Clips longer than the configured maximum duration (approximately six "
    "seconds) are truncated, and the WavLM feature extractor normalises the waveform and produces "
    "an attention mask that marks valid versus padded frames. Two lightweight quality gates reject "
    "unusable input: a root-mean-square energy below 0.001 is reported as silence, and a duration "
    "below 0.5 s is reported as too short, so that empty or near-empty recordings are surfaced "
    "to the user rather than scored.", indent=False)
h2(doc, "B", "Speech Mood Model")
body(doc,
    "The core estimator is a custom multi-task network, illustrated in Fig. 2. It is deliberately "
    "not a standard audio-classification head; rather, it shares one acoustic representation "
    "across two related tasks. The components are:", indent=False)
bullet(doc, "WavLM backbone.",
    " a microsoft/wavlm-base-plus encoder (hidden size 768). Following standard practice the "
    "convolutional feature encoder is frozen and the transformer layers are fine-tuned with a "
    "lower learning rate than the heads.")
bullet(doc, "Attentive masked pooling.",
    " a learned attention layer aggregates the variable-length frame sequence "
    "(B×T×768) into a single utterance embedding (B×768), respecting the padding mask "
    "so that padded frames receive zero weight.")
bullet(doc, "Mood head.",
    " a two-layer MLP (Linear–GELU–Dropout–Linear) producing six mood logits "
    "for Calm, Neutral, Content, Anxious, Agitated, and Low.")
bullet(doc, "Arousal head.",
    " a single linear layer producing two logits for low versus high arousal. The arousal target "
    "for each example is derived deterministically from its mood label (Calm/Neutral/Low → low; "
    "Content/Anxious/Agitated → high).")
bullet(doc, "Calibration.",
    " a single scalar temperature T, fit on the validation set after training and stored with the "
    "checkpoint, divides the logits before the softmax so that the reported probabilities are "
    "calibrated.")
add_figure(doc, F_MODEL,
    ("Fig. 2.", "Multi-task speech mood model. A frozen WavLM-Base-Plus feature encoder with "
     "fine-tuned transformer layers produces frame features that are aggregated by attentive "
     "masked pooling into a single embedding, which feeds a six-class mood head and a binary "
     "arousal head. Logits are divided by a fitted temperature (T ≈ 1.03) before the softmax."),
    width=2.7)
body(doc,
    "At inference, the calibrated softmax yields a probability distribution over the six moods and "
    "over the two arousal levels. The service returns the top-1 mood, its confidence, the full "
    "top-k distribution, the arousal-head prediction, and the arousal implied by the top mood as a "
    "cross-check. Two optional variance-reduction mechanisms are supported: test-time augmentation "
    "(TTA), which averages predictions over several augmented views of the same clip, and "
    "multi-seed ensembling, which averages calibrated probabilities across checkpoints. Both are "
    "disabled by default in the deployed service to minimise latency. An abstention threshold can "
    "be applied so that low-confidence clips are labelled uncertain and escalated to a human "
    "instead of acted upon.")
h2(doc, "C", "Training Strategy")
body(doc,
    "The model is trained on four merged public acted-emotion corpora—RAVDESS, CREMA-D, TESS, "
    "and SAVEE—whose original emotion labels are remapped to the six clinical mood classes "
    "(for example, happy→Content, fearful→Anxious, angry/disgust→Agitated, sad→Low). "
    "Splits are speaker-disjoint and corpus-aware: no speaker appears in more than one split, and "
    "the validation set is sized realistically at roughly 15%. This protocol prevents the speaker "
    "leakage that would otherwise inflate accuracy. Table II summarises the architecture and "
    "training configuration.", indent=False)
body(doc, "The training objective and regularisers target imbalance, calibration, and robustness:")
bullet(doc, "Class-balanced focal loss.",
    " effective-number class weighting combined with focal loss handles the heavy imbalance "
    "(Calm is a small minority) without the degenerate over-prediction of naive inverse-frequency "
    "weighting.")
bullet(doc, "Auxiliary arousal task.",
    " the arousal head is trained jointly as a weighted auxiliary loss, regularising the shared "
    "representation and yielding the reliable coarse signal.")
bullet(doc, "Domain-aware augmentation.",
    " mild waveform augmentation (additive noise, gain, ≤2-semitone pitch shift, time-stretch) "
    "plus telephony band-pass (300–3400 Hz) and reverberation to simulate real microphones "
    "and channels; intensity is ramped over early epochs. SpecAugment masking is applied on "
    "hidden states during training only.")
bullet(doc, "Optimisation.",
    " discriminative learning rates, a cosine schedule with warm-up, progressive unfreezing of "
    "the encoder, gradient clipping, mixed precision, and early stopping on validation macro-F1. "
    "Model selection and calibration use the validation set only; the test set is read once.")

add_table(doc, "II", "Mood Model Architecture and Training Configuration",
    ["Item", "Value"],
    [["Backbone", "WavLM-Base-Plus (hidden 768)"],
     ["Feature encoder", "Frozen (CNN); transformer fine-tuned"],
     ["Pooling", "Attentive masked pooling"],
     ["Mood head", "Linear-GELU-Dropout-Linear → 6"],
     ["Arousal head", "Linear → 2 (low / high)"],
     ["Sample rate", "16 kHz mono, ≤6 s"],
     ["Loss", "Class-balanced focal + aux. arousal"],
     ["Augmentation", "Noise, gain, pitch, telephony, reverb"],
     ["Calibration", "Temperature scaling, T ≈ 1.03"],
     ["Selection", "Early stop on val. macro-F1"]],
    col_w=[1.15, 2.1])

h2(doc, "D", "RAG Clinical Assistant")
body(doc,
    "The assistant, shown in Fig. 3, answers clinical and caregiving questions while guaranteeing "
    "that patient-specific statements are grounded in the database. Incoming questions are first "
    "classified by a semantic router: the question is embedded with a Sentence-Transformers "
    "all-MiniLM-L6-v2 encoder (384 dimensions) and compared by cosine similarity against three "
    "intent prototypes—patient, knowledge, and hybrid. If the two top intents are within 0.05 "
    "of each other the query is treated as hybrid; if the best similarity is below 0.65 the router "
    "defers to the LLM for a single-word classification; otherwise the highest-scoring intent is "
    "used.", indent=False)
add_figure(doc, F_RAG,
    ("Fig. 3.", "Clinically-safe RAG pipeline. A MiniLM semantic router classifies each question "
     "as patient, knowledge, or hybrid. Patient facts are read verbatim from the MongoDB record "
     "(the sole source of truth); general medical context is retrieved from a FAISS knowledge "
     "base of curated PDFs and web sources. A guarded prompt is generated for a deterministic "
     "Groq-hosted Llama-3.3-70B model, and every answer carries a safety disclaimer."))
body(doc,
    "General medical context is retrieved from a FAISS vector store built from approximately "
    "ninety Alzheimer’s and dementia PDFs (organised across eight topical folders) together with "
    "a small number of authoritative web sources, chunked with a recursive character splitter "
    "(500-character chunks, 100-character overlap) and embedded with the same MiniLM model. "
    "Retrieval fetches the six nearest chunks, keeps those below an L2 distance threshold of 1.20 "
    "(with a floor of two results), and returns at most four chunks as context.")
body(doc,
    "Generation uses a Groq-hosted Llama-3.3-70B model at temperature 0 for deterministic, "
    "reproducible output. Three modes are supported. In patient mode (clinician audience) and "
    "family mode (compassionate, jargon-free audience) the prompt embeds the MongoDB patient "
    "record as the source of truth; in general mode, with no patient selected, only the knowledge "
    "base is used. A set of explicit anti-hallucination rules is enforced in the prompt: every "
    "patient claim must be traceable to a database field, the recorded Alzheimer’s stage must be "
    "reproduced verbatim, missing fields must be reported as “not recorded in the patient’s "
    "file,” retrieved knowledge may enrich but never override a database fact, and user claims "
    "that contradict the record are politely corrected. A short conversational memory (the last "
    "ten exchanges, isolated per user-and-patient session) provides continuity, and a medical "
    "disclaimer is appended to every response. Table V lists the configuration.")

# --------------------------------------------------------------------------- V. EVAL
h1(doc, "V", "Experiments and Evaluation")
body(doc,
    "We evaluate the two subsystems independently. The mood model is assessed on a speaker-disjoint "
    "test split of the merged corpora using accuracy, macro-F1, per-class confusion, training "
    "dynamics, and calibration error. The RAG assistant is assessed on a held-out set of twelve "
    "clinical/caregiving queries using standard retrieval metrics and an embedding-based semantic "
    "similarity between generated answers and reference answers.", indent=False)

h2(doc, "A", "Mood Model: Setup and Metrics")
body(doc,
    "Reported figures span the range observed with and without TTA and ensembling. Six-class "
    "accuracy and macro-F1 measure fine mood recognition; low/high arousal accuracy measures the "
    "coarser, clinically prioritised signal; expected calibration error (ECE) measures how well "
    "the reported confidences match empirical correctness. Macro-F1 is the early-stopping "
    "criterion, which is appropriate given the class imbalance.", indent=False)

add_table(doc, "III", "Speech Mood Model Performance",
    ["Metric", "Result"],
    [["6-class accuracy", "0.74 – 0.77"],
     ["6-class macro-F1", "0.73 – 0.76"],
     ["Arousal accuracy (low/high)", "0.86 – 0.88"],
     ["Peak validation macro-F1", "0.82 (epoch 9)"],
     ["ECE before calibration", "0.041"],
     ["ECE after calibration (T≈1.03)", "0.039"]],
    col_w=[1.85, 1.4])

body(doc,
    "Fig. 4 plots validation macro-F1 over training for the primary seed. Performance rises "
    "steeply in the first epochs and plateaus around 0.79–0.82, peaking at epoch 9; the chosen "
    "checkpoint achieves a validation macro-F1 of 0.816. The early plateau is consistent with the "
    "label-ambiguity ceiling of acted emotion data rather than with under-training.")
add_figure(doc, M2,
    ("Fig. 4.", "Validation macro-F1 across training epochs (seed 42). The model converges quickly "
     "and peaks near epoch 9 at a macro-F1 of approximately 0.82."))
body(doc,
    "Fig. 5 shows the row-normalised confusion matrix on the test split. The diagonal "
    "recalls—Calm 0.97, Neutral 0.87, Agitated 0.78, Content 0.68, Low 0.68, and Anxious "
    "0.66—reveal the structure of the residual error. The dominant confusions are "
    "within-arousal: Content is most often confused with the other high-arousal classes, and the "
    "low-to-moderate-arousal negative states (Anxious, Low) are confused with their neighbours. "
    "This pattern confirms that fine mood error concentrates where valence is acoustically "
    "under-determined, and that collapsing to arousal recovers accuracy. The exceptionally high "
    "Calm recall partly reflects a corpus confound—Calm originates from a single corpus—and "
    "should be read with caution.")
add_figure(doc, M1,
    ("Fig. 5.", "Row-normalised confusion matrix on the speaker-disjoint test split. Errors "
     "concentrate within the same arousal band, e.g. Content↔Agitated and Anxious↔Low."))
body(doc,
    "Fig. 6 presents reliability diagrams before and after temperature scaling. The model is "
    "already well calibrated—the curves track the diagonal closely—and a fitted temperature "
    "of 1.03 lowers the expected calibration error marginally from 0.041 to 0.039. The near-unity "
    "temperature indicates that the class-balanced focal objective produces confidences that are "
    "intrinsically reliable, which is important because downstream alerting and abstention act on "
    "these probabilities.")
add_figure(doc, M3,
    ("Fig. 6.", "Reliability diagrams before (left) and after (right) temperature scaling. "
     "Calibration is good a priori; the fitted temperature T = 1.03 reduces ECE from 0.041 to "
     "0.039."))

h2(doc, "B", "RAG Assistant: Setup and Metrics")
body(doc,
    "The RAG assistant is evaluated on twelve held-out queries spanning patient-specific, "
    "general-knowledge, and hybrid intents. Retrieval quality is measured by precision (the "
    "fraction of retrieved chunks that are relevant), mean reciprocal rank (MRR, the rank of the "
    "first relevant chunk), and normalized discounted cumulative gain (NDCG, which rewards placing "
    "relevant chunks early). Routing accuracy measures whether the semantic router assigns the "
    "correct intent. Answer quality is measured by the cosine semantic similarity between the "
    "generated answer and a reference answer in the MiniLM embedding space. Table VI summarises "
    "the results.", indent=False)

add_table(doc, "IV", "RAG Retrieval and Generation Results (n = 12)",
    ["Metric", "Score"],
    [["Retrieval precision", "0.95"],
     ["Mean reciprocal rank (MRR)", "0.95"],
     ["NDCG", "0.97"],
     ["Routing accuracy", "1.00"],
     ["Mean answer semantic similarity", "0.96"]],
    col_w=[1.95, 1.3])

body(doc,
    "Fig. 7 summarises the four retrieval/routing metrics as a bar chart, and Fig. 8 presents the "
    "same scores on a radar plot, where the near-maximal enclosed area visualises uniformly strong "
    "performance across precision, MRR, NDCG, and routing. Routing accuracy reaches 1.00 on this "
    "set, indicating that the lightweight cosine-similarity router correctly separates "
    "patient-specific from general-knowledge intents.")
add_figure(doc, R2,
    ("Fig. 7.", "Retrieval evaluation summary: precision, MRR, NDCG, and routing accuracy, all at "
     "or above 0.95."))
add_figure(doc, R5,
    ("Fig. 8.", "Radar view of the four retrieval metrics; the near-maximal area indicates "
     "uniformly strong retrieval and routing."))
body(doc,
    "Figs. 9–11 examine answer–reference semantic similarity in detail. The per-sample scatter "
    "(Fig. 9) shows that every one of the twelve queries scores between 0.82 and 1.00, with most "
    "near 1.0; the distribution (Fig. 10) is tightly concentrated with a median near 0.96; and the "
    "cumulative curve (Fig. 11) confirms that the bulk of probability mass lies above a 0.96 "
    "similarity threshold. Together these indicate that the grounded generator reproduces "
    "reference answers faithfully, with no catastrophic failures on the evaluated queries.")
add_figure(doc, R3,
    ("Fig. 9.", "Per-sample answer–reference semantic similarity for the twelve evaluation "
     "queries; all scores fall between 0.82 and 1.00."))
add_figure(doc, R4,
    ("Fig. 10.", "Distribution of semantic-similarity scores (violin plot), concentrated near a "
     "median of approximately 0.96."))
add_figure(doc, R1,
    ("Fig. 11.", "Cumulative semantic-similarity curve: the proportion of samples meeting or "
     "exceeding each similarity threshold."))
body(doc,
    "Finally, Fig. 12 plots the cumulative (running-average) retrieval metrics as queries are "
    "added. After an initial transient over the first few queries, precision, MRR, and NDCG all "
    "stabilise above 0.94, indicating that the reported averages are not driven by a small number "
    "of favourable cases but reflect consistent behaviour across the query set.")
add_figure(doc, R6,
    ("Fig. 12.", "Cumulative (running-average) precision, MRR, and NDCG. All three stabilise "
     "above 0.94 after the first few queries."))

# --------------------------------------------------------------------------- VI. DISCUSSION
h1(doc, "VI", "Discussion")
h2(doc, "A", "Strengths")
body(doc,
    "The system’s principal strength is that it turns an otherwise unobserved signal—a "
    "patient’s moment-to-moment affective state—into a calibrated, queryable, and shareable "
    "record without specialised hardware. The multi-task design is well matched to the problem: by "
    "predicting arousal jointly with mood, the model exposes a coarse signal (0.86–0.88 "
    "accuracy) that is substantially more reliable than fine six-class mood and that aligns with "
    "the clinical priority of detecting agitation or distress. Calibrated probabilities and an "
    "abstention mechanism make the output safe to act on, and the strict grounding rules in the "
    "RAG assistant directly target the dominant risk of clinical LLMs—fabricated patient facts. "
    "Architecturally, decoupling inference into Python sidecars keeps the model independent of the "
    "web stack and allows the acoustic model to be upgraded without touching the application.",
    indent=False)
h2(doc, "B", "Limitations")
body(doc,
    "Several limitations bound the interpretation of our results. First, the mood model is trained "
    "and evaluated on acted emotional speech recorded in clean conditions, mostly from younger "
    "speakers; real spontaneous speech from elderly patients in noisier settings constitutes an "
    "unmeasured domain shift, and real-world accuracy should be expected to be lower than the "
    "reported in-distribution figures. Second, a meaningful share of the six-class error is "
    "irreducible label ambiguity in perceived-emotion data, and some per-class results (notably "
    "Calm) are corpus-confounded. Third, the RAG evaluation set is small (twelve queries); the "
    "near-perfect retrieval and routing scores are encouraging but warrant a larger, "
    "independently-labelled benchmark before strong claims are made, and the similarity metric "
    "measures agreement with reference answers rather than clinical correctness per se. Finally, "
    "the system is a research prototype: it is not a medical device, has not undergone prospective "
    "clinical validation, and must not be used for autonomous clinical decisions.", indent=False)
h2(doc, "C", "Interpretation")
body(doc,
    "Read together, the results support a clear design stance: treat fine mood as a soft, "
    "uncertain hint, treat arousal as the dependable monitoring signal, act on calibrated "
    "probabilities with abstention, and aggregate over time rather than over-interpreting any "
    "single clip. The RAG assistant should likewise be understood as a grounded information layer "
    "that surfaces and explains recorded facts and curated knowledge, not as an autonomous "
    "reasoner over patient care.", indent=False)

# --------------------------------------------------------------------------- VII. CONCLUSION
h1(doc, "VII", "Conclusion and Future Work")
body(doc,
    "We presented ALZCare, an integrated platform that estimates patient mood and arousal from "
    "short voice recordings and delivers grounded clinical assistance through a retrieval-augmented "
    "language model. The speech model—a calibrated WavLM multi-task network—achieves "
    "0.74–0.77 six-class accuracy, 0.73–0.76 macro-F1, and 0.86–0.88 arousal accuracy, "
    "while the RAG assistant attains 0.95 retrieval precision, 0.97 NDCG, perfect routing, and "
    "0.96 mean answer similarity on a held-out query set. The whole system is deployed as a "
    "role-aware MERN application with real-time distress alerting.", indent=False)
body(doc,
    "Future work follows directly from the limitations. The highest-value next step is domain "
    "adaptation to clinician-labelled spontaneous patient speech, accompanied by honest "
    "leave-one-corpus-out and prospective evaluation. A dedicated, calibrated binary "
    "distress/agitation detector may prove both more reliable and more clinically useful than fine "
    "mood classification. Temporal modelling—aggregating predictions across a monitoring window "
    "with uncertainty propagation—would convert noisy per-clip estimates into trustworthy "
    "trends. On the assistant side, a larger independently-labelled evaluation benchmark, citation "
    "of retrieved passages in answers, and subgroup fairness analysis are the natural priorities. "
    "Scalability and real-world deployment would further benefit from containerised, horizontally "
    "scalable inference services and on-device pre-processing to reduce the transmission of raw "
    "audio.")

# Table V - RAG config (placed near references for completeness)
add_table(doc, "V", "RAG Assistant Configuration",
    ["Item", "Value"],
    [["Embedding model", "all-MiniLM-L6-v2 (384-d)"],
     ["Vector store", "FAISS (local)"],
     ["Chunking", "500 chars, 100 overlap"],
     ["Retrieval", "k = 6, L2 < 1.20, ≤4 kept"],
     ["Knowledge base", "≈90 PDFs + web sources"],
     ["Generator", "Groq Llama-3.3-70B (T = 0)"],
     ["Router thresholds", "conf. 0.65, gap 0.05"],
     ["Memory", "last 10 exchanges / session"]],
    col_w=[1.2, 2.05])

# --------------------------------------------------------------------------- REFERENCES
h1(doc, "", "References")
# fix: References heading should not have a leading dot
refs = [
    "A. Baevski, H. Zhou, A. Mohamed, and M. Auli, “wav2vec 2.0: A framework for "
    "self-supervised learning of speech representations,” in Adv. Neural Inf. Process. Syst. "
    "(NeurIPS), vol. 33, 2020, pp. 12449–12460.",
    "W.-N. Hsu, B. Bolte, Y.-H. H. Tsai, K. Lakhotia, R. Salakhutdinov, and A. Mohamed, "
    "“HuBERT: Self-supervised speech representation learning by masked prediction of hidden "
    "units,” IEEE/ACM Trans. Audio, Speech, Lang. Process., vol. 29, pp. 3451–3460, 2021.",
    "S. Chen et al., “WavLM: Large-scale self-supervised pre-training for full stack speech "
    "processing,” IEEE J. Sel. Topics Signal Process., vol. 16, no. 6, pp. 1505–1518, 2022.",
    "S. Yang et al., “SUPERB: Speech processing universal performance benchmark,” in Proc. "
    "Interspeech, 2021, pp. 1194–1198.",
    "S. R. Livingstone and F. A. Russo, “The Ryerson audio-visual database of emotional speech "
    "and song (RAVDESS),” PLoS ONE, vol. 13, no. 5, e0196391, 2018.",
    "H. Cao, D. G. Cooper, M. K. Keutmann, R. C. Gur, A. Nenkova, and R. Verma, “CREMA-D: "
    "Crowd-sourced emotional multimodal actors dataset,” IEEE Trans. Affect. Comput., vol. 5, "
    "no. 4, pp. 377–390, 2014.",
    "Y. Cui, M. Jia, T.-Y. Lin, Y. Song, and S. Belongie, “Class-balanced loss based on "
    "effective number of samples,” in Proc. IEEE/CVF CVPR, 2019, pp. 9268–9277.",
    "T.-Y. Lin, P. Goyal, R. Girshick, K. He, and P. Dollár, “Focal loss for dense object "
    "detection,” in Proc. IEEE Int. Conf. Comput. Vis. (ICCV), 2017, pp. 2980–2988.",
    "D. S. Park et al., “SpecAugment: A simple data augmentation method for automatic speech "
    "recognition,” in Proc. Interspeech, 2019, pp. 2613–2617.",
    "C. Guo, G. Pleiss, Y. Sun, and K. Q. Weinberger, “On calibration of modern neural "
    "networks,” in Proc. Int. Conf. Mach. Learn. (ICML), 2017, pp. 1321–1330.",
    "P. Lewis et al., “Retrieval-augmented generation for knowledge-intensive NLP tasks,” in "
    "Adv. Neural Inf. Process. Syst. (NeurIPS), vol. 33, 2020, pp. 9459–9474.",
    "N. Reimers and I. Gurevych, “Sentence-BERT: Sentence embeddings using Siamese "
    "BERT-networks,” in Proc. Conf. Empirical Methods Nat. Lang. Process. (EMNLP), 2019, pp. "
    "3982–3992.",
    "W. Wang, F. Wei, L. Dong, H. Bao, N. Yang, and M. Zhou, “MiniLM: Deep self-attention "
    "distillation for task-agnostic compression of pre-trained transformers,” in Adv. Neural "
    "Inf. Process. Syst. (NeurIPS), vol. 33, 2020, pp. 5776–5788.",
    "J. Johnson, M. Douze, and H. Jégou, “Billion-scale similarity search with GPUs,” "
    "IEEE Trans. Big Data, vol. 7, no. 3, pp. 535–547, 2021.",
    "A. Dubey et al., “The Llama 3 herd of models,” arXiv preprint arXiv:2407.21783, 2024.",
    "K. He, X. Zhang, S. Ren, and J. Sun, “Deep residual learning for image recognition,” in "
    "Proc. IEEE/CVF CVPR, 2016, pp. 770–778.",
    "P. Khosla et al., “Supervised contrastive learning,” in Adv. Neural Inf. Process. Syst. "
    "(NeurIPS), vol. 33, 2020, pp. 18661–18673.",
    "A. Vaswani et al., “Attention is all you need,” in Adv. Neural Inf. Process. Syst. "
    "(NeurIPS), vol. 30, 2017, pp. 5998–6008.",
    "K. Järvelin and J. Kekäläinen, “Cumulated gain-based evaluation of IR techniques,” "
    "ACM Trans. Inf. Syst., vol. 20, no. 4, pp. 422–446, 2002.",
    "G. M. McKhann et al., “The diagnosis of dementia due to Alzheimer’s disease,” "
    "Alzheimer’s & Dementia, vol. 7, no. 3, pp. 263–269, 2011.",
    "C. Ballard, J. Corbett, and S. Reichelt, “Management of neuropsychiatric symptoms in "
    "people with dementia,” CNS Drugs, vol. 25, no. 9, pp. 729–739, 2011.",
    "J. Devlin, M.-W. Chang, K. Lee, and K. Toutanova, “BERT: Pre-training of deep "
    "bidirectional transformers for language understanding,” in Proc. NAACL-HLT, 2019, pp. "
    "4171–4186.",
]
for i, r in enumerate(refs, 1):
    add_ref(doc, i, r)

OUT_DOCX = os.path.join(OUT, "ALZCare_IEEE_Paper.docx")
doc.save(OUT_DOCX)
print("Saved:", OUT_DOCX)
